"""Extracción de texto de archivos subidos (PDF, Word .docx, texto/transcripciones)."""
import io


class UnsupportedFile(Exception):
    """El archivo no es de un tipo del que sepamos extraer texto."""


_PDF = "application/pdf"
_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_TEXT_EXTS = (".txt", ".md", ".csv", ".vtt", ".srt", ".log")


def extract_text(file_name: str, mime_type: str | None, data: bytes) -> str:
    name = (file_name or "").lower()
    mime = (mime_type or "").lower()

    if mime == _PDF or name.endswith(".pdf"):
        return _from_pdf(data)
    if mime == _DOCX or name.endswith(".docx"):
        return _from_docx(data)
    if mime.startswith("text/") or name.endswith(_TEXT_EXTS):
        return data.decode("utf-8", errors="ignore").strip()

    # Último intento: tratarlo como texto plano.
    try:
        return data.decode("utf-8").strip()
    except UnicodeDecodeError as exc:
        raise UnsupportedFile(f"Formato no soportado: {file_name}") from exc


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    return "\n\n".join(p for p in pages if p).strip()


def _from_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()
